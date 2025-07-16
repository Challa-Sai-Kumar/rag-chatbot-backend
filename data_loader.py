import os
import requests
from bs4 import BeautifulSoup
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from docx import Document as DocxDocument
import chromadb
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import time
from transformers import AutoTokenizer

class DataLoader:
    def __init__(self):
        # Tokenizer-aware text splitter
        tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=300,
            chunk_overlap=50,
            length_function=lambda x: len(tokenizer.tokenize(x)),
        )
        self.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        self.chroma_path = "chroma_db"

    def clean_text(self, text):
        return text.replace("\xa0", " ").replace("\n", " ").strip()

    def scrape_angelone_support(self):
        """Scrape support pages from angelone.in"""
        base_url = "https://www.angelone.in/support"
        documents = []
        visited = set()
        to_visit = [base_url]

        while to_visit and len(visited) < 30:
            url = to_visit.pop(0)
            if url in visited:
                continue

            visited.add(url)
            try:
                response = requests.get(url, timeout=10)
                soup = BeautifulSoup(response.content, 'html.parser')

                main_content = soup.find('main') or soup.find('div', class_='content')
                if main_content:
                    text = self.clean_text(main_content.get_text())
                    if text and len(text) > 100:
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source": url,
                                "type": "webpage",
                                "scraped_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                                "length": len(text)
                            }
                        ))

                for link in soup.find_all('a', href=True):
                    href = link['href']
                    if "/support" in href:
                        full_url = href if href.startswith("http") else f"https://www.angelone.in{href}"
                        if full_url not in visited:
                            to_visit.append(full_url)

                time.sleep(1)  # Respectful crawling

            except Exception as e:
                print(f"Error scraping {url}: {e}")
                continue

        return documents

    def load_pdf(self, pdf_path):
        loader = PyPDFLoader(pdf_path)
        documents = loader.load()
        return documents

    def load_docx(self, docx_path):
        doc = DocxDocument(docx_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        cleaned_text = self.clean_text(text)
        return [Document(
            page_content=cleaned_text,
            metadata={"source": docx_path, "type": "docx", "length": len(cleaned_text)}
        )]

    def load_documents_from_folder(self, folder_path):
        documents = []
        if os.path.exists(folder_path):
            for filename in os.listdir(folder_path):
                file_path = os.path.join(folder_path, filename)
                try:
                    if filename.endswith('.pdf'):
                        documents.extend(self.load_pdf(file_path))
                    elif filename.endswith('.docx'):
                        documents.extend(self.load_docx(file_path))
                except Exception as e:
                    print(f"Error loading {filename}: {e}")
        return documents

    def create_database(self):
        print("Loading documents...")
        web_documents = self.scrape_angelone_support()
        print(f"Loaded {len(web_documents)} web documents")

        local_documents = self.load_documents_from_folder("documents")
        print(f"Loaded {len(local_documents)} local documents")

        all_documents = web_documents + local_documents
        if not all_documents:
            print("No documents found!")
            return

        print("Splitting documents...")
        chunks = self.text_splitter.split_documents(all_documents)
        print(f"Created {len(chunks)} chunks")

        if os.path.exists(self.chroma_path):
            import shutil
            shutil.rmtree(self.chroma_path)

        print("Creating vector database...")
        vectorstore = Chroma.from_documents(
            chunks,
            self.embeddings,
            persist_directory=self.chroma_path
        )

        print(f"Database created successfully with {len(chunks)} chunks!")
        return vectorstore

if __name__ == "__main__":
    loader = DataLoader()
    loader.create_database()
